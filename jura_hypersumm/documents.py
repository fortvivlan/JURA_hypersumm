"""Court-decision parsing helpers."""

from __future__ import annotations

import re
from pathlib import Path

POSTANOVIL_PATTERN = re.compile(
    r"П\s*О\s*С\s*Т\s*А\s*Н\s*О\s*В\s*И\s*Л\s*[:.]?",
    flags=re.IGNORECASE,
)
IRRELEVANT_SENTENCE_MARKERS = (
    "реквизит",
    "ре...изит",
    "квитанци",
    "судья",
)
BANK_DETAIL_PATTERN = re.compile(
    r"(?<!\w)(?:БИК|ИНН|УИН|КБК|ОКТМО|л/с|р/с)(?!\w)"
)
NUMERIC_SENTENCE_PATTERN = re.compile(r"(?:\d+\s*)+[.!?]?")
OTD_ABBREVIATION = "Отд."
OTD_PLACEHOLDER = "Отд\ue000"


def read_docx_text(path: str | Path) -> str:
    """Read paragraphs and table cells from a DOCX file."""
    from docx import Document

    document = Document(Path(path))
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            blocks.extend(cell.text for cell in row.cells if cell.text)
    return "\n".join(blocks)


def extract_operative_section(text: str) -> str | None:
    """Return text after the final ПОСТАНОВИЛ marker, or ``None`` if absent."""
    matches = list(POSTANOVIL_PATTERN.finditer(text))
    if not matches:
        return None
    section = text[matches[-1].end() :].strip()
    return section or None


def split_russian_sentences(text: str) -> list[str]:
    """Split sentences without treating the ``Отд.`` abbreviation as an ending."""
    from razdel import sentenize

    protected_text = text.replace(OTD_ABBREVIATION, OTD_PLACEHOLDER)
    return [
        sentence.text.replace(OTD_PLACEHOLDER, OTD_ABBREVIATION).strip()
        for sentence in sentenize(protected_text)
        if sentence.text.strip()
    ]


def textcheck(text: str) -> bool:
    """Return whether a sentence is irrelevant to codex contradiction checks."""
    normalized = text.lower()
    return any(
        marker in normalized for marker in IRRELEVANT_SENTENCE_MARKERS
    ) or bool(
        BANK_DETAIL_PATTERN.search(text)
        or NUMERIC_SENTENCE_PATTERN.fullmatch(text.strip())
    )
